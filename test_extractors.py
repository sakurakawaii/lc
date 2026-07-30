"""
Unit tests for extractors.py.

PDF reading is mocked via PyPDF2.PdfReader (constructing a realistic PDF
in-process isn't practical without a PDF-writing library). DOCX and EML tests
use real python-docx documents / real MIME messages built in-memory or in a
temp directory, since both are cheap and fast to construct for real. No
checked-in fixture files or network calls are needed to run these.
"""

import base64
import io
import os
import tempfile
import unittest
from email.message import EmailMessage
from unittest.mock import MagicMock, mock_open, patch

from docx import Document

from extractors import (
    extract_image,
    extract_text_from_docx,
    extract_text_from_eml,
    extract_text_from_pdf,
)

FAKE_IMAGE_BYTES = b"\x89PNG\r\n\x1a\nfake-binary-image-data"


class TestExtractTextFromPdf(unittest.TestCase):
    def _mock_reader(self, page_texts):
        pages = []
        for text in page_texts:
            page = MagicMock()
            if isinstance(text, Exception):
                page.extract_text.side_effect = text
            else:
                page.extract_text.return_value = text
            pages.append(page)
        reader = MagicMock()
        reader.pages = pages
        return reader

    @patch("extractors.PyPDF2.PdfReader")
    def test_concatenates_text_from_all_pages(self, mock_reader_cls):
        mock_reader_cls.return_value = self._mock_reader(["Page one.", "Page two."])

        with patch("builtins.open", mock_open(read_data=b"%PDF-fake")):
            text = extract_text_from_pdf("sample.pdf")

        self.assertEqual(text, "Page one.\nPage two.")

    @patch("extractors.PyPDF2.PdfReader")
    def test_a_failing_page_is_skipped_without_crashing(self, mock_reader_cls):
        mock_reader_cls.return_value = self._mock_reader(
            ["Good page.", RuntimeError("corrupt page")]
        )

        with patch("builtins.open", mock_open(read_data=b"%PDF-fake")):
            text = extract_text_from_pdf("sample.pdf")

        self.assertEqual(text, "Good page.")

    def test_missing_file_returns_empty_string(self):
        with patch("builtins.open", side_effect=FileNotFoundError("no such file")):
            text = extract_text_from_pdf("missing.pdf")

        self.assertEqual(text, "")


class TestExtractTextFromDocx(unittest.TestCase):
    def test_extracts_and_joins_non_empty_paragraphs(self):
        document = Document()
        document.add_paragraph("First paragraph.")
        document.add_paragraph("")
        document.add_paragraph("Second paragraph.")
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        text = extract_text_from_docx(buffer)

        self.assertEqual(text, "First paragraph.\nSecond paragraph.")

    @patch("extractors.Document", side_effect=RuntimeError("corrupt docx"))
    def test_unreadable_file_returns_empty_string(self, mock_document_cls):
        text = extract_text_from_docx("broken.docx")

        self.assertEqual(text, "")


class TestExtractImage(unittest.TestCase):
    def test_encodes_png_file_and_returns_correct_media_type(self):
        with patch("builtins.open", mock_open(read_data=FAKE_IMAGE_BYTES)):
            encoded, media_type = extract_image("photo.png")

        self.assertEqual(media_type, "image/png")
        self.assertEqual(encoded, base64.b64encode(FAKE_IMAGE_BYTES).decode("ascii"))

    def test_encodes_jpg_file_and_returns_correct_media_type(self):
        with patch("builtins.open", mock_open(read_data=FAKE_IMAGE_BYTES)):
            encoded, media_type = extract_image("photo.jpg")

        self.assertEqual(media_type, "image/jpeg")
        self.assertEqual(encoded, base64.b64encode(FAKE_IMAGE_BYTES).decode("ascii"))

    def test_encodes_jpeg_file_and_returns_correct_media_type(self):
        with patch("builtins.open", mock_open(read_data=FAKE_IMAGE_BYTES)):
            encoded, media_type = extract_image("photo.jpeg")

        self.assertEqual(media_type, "image/jpeg")
        self.assertEqual(encoded, base64.b64encode(FAKE_IMAGE_BYTES).decode("ascii"))

    def test_case_insensitive_extension_matching(self):
        with patch("builtins.open", mock_open(read_data=FAKE_IMAGE_BYTES)):
            encoded, media_type = extract_image("PHOTO.PNG")

        self.assertEqual(media_type, "image/png")
        self.assertIsNotNone(encoded)

    def test_unsupported_extension_returns_none_without_opening_file(self):
        with patch("builtins.open", mock_open(read_data=FAKE_IMAGE_BYTES)) as mock_open_file:
            encoded, media_type = extract_image("document.gif")

        self.assertIsNone(encoded)
        self.assertIsNone(media_type)
        mock_open_file.assert_not_called()

    def test_missing_file_returns_none_none(self):
        with patch("builtins.open", side_effect=FileNotFoundError("no such file")):
            encoded, media_type = extract_image("missing.png")

        self.assertIsNone(encoded)
        self.assertIsNone(media_type)


class TestExtractTextFromEml(unittest.TestCase):
    """
    Builds real MIME messages in a temp directory rather than mocking
    open()/Path, since extract_text_from_eml's attachment-saving path
    exercises os.makedirs/pathlib.Path/file writes together and mocking all
    of that convincingly would be more fragile than just using the real
    filesystem in a directory that's discarded at the end of each test.
    """

    def _build_eml_bytes(self, body="Hello, this is the body.", attachment=None):
        msg = EmailMessage()
        msg["Subject"] = "Test email"
        msg["From"] = "sender@example.com"
        msg["To"] = "recipient@example.com"
        msg.set_content(body)
        if attachment is not None:
            filename, data = attachment
            msg.add_attachment(
                data, maintype="application", subtype="octet-stream", filename=filename
            )
        return msg.as_bytes()

    def test_extracts_plain_text_body_with_no_attachments(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            eml_path = os.path.join(tmp_dir, "email.eml")
            with open(eml_path, "wb") as f:
                f.write(self._build_eml_bytes(body="Hello, this is the body."))

            body_text, attachment_paths = extract_text_from_eml(
                eml_path, attachment_dir=os.path.join(tmp_dir, "attachments")
            )

        self.assertIn("Hello, this is the body.", body_text)
        self.assertEqual(attachment_paths, [])

    def test_saves_attachment_and_reports_its_readable_path(self):
        with tempfile.TemporaryDirectory() as tmp_dir:
            eml_path = os.path.join(tmp_dir, "email.eml")
            with open(eml_path, "wb") as f:
                f.write(self._build_eml_bytes(attachment=("notes.txt", b"attachment contents")))
            attachment_dir = os.path.join(tmp_dir, "attachments")

            body_text, attachment_paths = extract_text_from_eml(
                eml_path, attachment_dir=attachment_dir
            )

            self.assertEqual(len(attachment_paths), 1)
            saved_path = attachment_paths[0]
            self.assertTrue(saved_path.endswith("notes.txt"))
            with open(saved_path, "rb") as f:
                self.assertEqual(f.read(), b"attachment contents")

    def test_unreadable_file_returns_empty_body_and_no_attachments(self):
        body_text, attachment_paths = extract_text_from_eml("/nonexistent/missing.eml")

        self.assertEqual(body_text, "")
        self.assertEqual(attachment_paths, [])


if __name__ == "__main__":
    unittest.main()
